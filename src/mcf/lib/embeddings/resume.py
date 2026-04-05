"""Resume text extraction helpers."""

from __future__ import annotations

import io
import re
from pathlib import Path


def extract_resume_text(source: str | Path | bytes) -> str:
    """Extract plain text from a resume file.

    *source* may be:
      - a file path (str or Path) — used by the local CLI flow
      - raw bytes — used by the upload endpoint (no temp file needed)

    Supported formats: .pdf, .docx, .txt, .md
    When *source* is bytes the format is detected by sniffing magic bytes
    (PDF starts with ``%PDF``; DOCX is a ZIP).
    """
    if isinstance(source, (str, Path)):
        return _extract_from_path(Path(source))
    return _extract_from_bytes(source)


def _extract_from_path(p: Path) -> str:
    suffix = p.suffix.lower()

    if suffix in {".txt", ".md"}:
        return p.read_text(encoding="utf-8", errors="ignore")

    if suffix == ".pdf":
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(p))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if suffix == ".docx":
        from docx import Document  # type: ignore

        doc = Document(str(p))
        return "\n".join(par.text for par in doc.paragraphs if par.text)

    raise ValueError(f"Unsupported resume file type: {suffix} (supported: .txt, .md, .pdf, .docx)")


def _extract_from_bytes(data: bytes) -> str:
    """Detect format from magic bytes and extract text."""
    if data[:4] == b"%PDF":
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    # DOCX is a ZIP file starting with PK\x03\x04
    if data[:2] == b"PK":
        from docx import Document  # type: ignore

        doc = Document(io.BytesIO(data))
        return "\n".join(par.text for par in doc.paragraphs if par.text)

    # Fall back: treat as plain text
    return data.decode("utf-8", errors="ignore")


# Section header patterns (case-insensitive, allow optional colon)
_SECTION_PATTERNS = [
    (r"^(?:technical\s+)?skills(?:\s+&\s+technologies)?\s*:?\s*$", "skills"),
    (r"^(?:core\s+)?competencies\s*:?\s*$", "skills"),
    (r"^(?:professional\s+)?summary\s*:?\s*$", "summary"),
    (r"^(?:career\s+)?objective\s*:?\s*$", "summary"),
    (r"^(?:executive\s+)?summary\s*:?\s*$", "summary"),
    (r"^experience\s*:?\s*$", "experience"),
    (r"^work\s+(?:history|experience)\s*:?\s*$", "experience"),
    (r"^employment\s+(?:history|experience)\s*:?\s*$", "experience"),
    (r"^professional\s+experience\s*:?\s*$", "experience"),
    (r"^projects?\s*:?\s*$", "projects"),
    (r"^education\s*:?\s*$", "education"),
    (r"^(?:academic\s+)?qualifications?\s*:?\s*$", "education"),
    (r"^certifications?\s*:?\s*$", "certifications"),
    (r"^languages?\s*:?\s*$", "languages"),
]


def preprocess_resume_text(raw_text: str) -> str:
    """Clean and reorder resume text for better embedding.

    Strips noise (emails, phones, URLs, addresses), detects section headers,
    and reorders sections so Skills and Experience come first. This maximizes
    the 512-token BGE signal budget for matching.
    """
    if not raw_text or not raw_text.strip():
        return raw_text

    # Strip noise lines
    lines = raw_text.splitlines()
    cleaned: list[str] = []
    for line in lines:
        s = line.strip()
        if not s:
            continue
        # Skip emails
        if re.match(r"^[\w\.-]+@[\w\.-]+\.\w+$", s):
            continue
        # Skip phone numbers (various formats)
        if re.match(r"^[\d\s\-\(\)\+\.]{10,}$", s):
            continue
        # Skip URLs
        if re.match(r"^https?://", s, re.I):
            continue
        if re.match(r"^www\.", s, re.I):
            continue
        if re.match(r"^linkedin\.com/in/", s, re.I):
            continue
        if re.match(r"^github\.com/", s, re.I):
            continue
        # Skip long separator lines
        if re.match(r"^[\-\=\*\.]{3,}$", s):
            continue
        # Skip page numbers
        if re.match(r"^page\s+\d+\s*$", s, re.I):
            continue
        if re.match(r"^\d+\s*$", s) and len(s) <= 3:
            continue
        # Skip lines that look like a full address (street + number pattern)
        if re.match(r"^\d+\s+[\w\s]+(?:street|st|avenue|ave|road|rd|blvd|drive|dr|lane|ln)\b", s, re.I):
            continue
        cleaned.append(s)

    # Rejoin and split into sections
    text = "\n".join(cleaned)
    sections: dict[str, str] = {}
    current_section = "preamble"
    current_content: list[str] = []

    def flush_section():
        if current_content:
            content = "\n".join(current_content).strip()
            if content:
                sections.setdefault(current_section, "")
                sections[current_section] = (
                    (sections[current_section] + "\n\n" + content).strip()
                    if sections[current_section]
                    else content
                )

    for line in text.splitlines():
        s = line.strip()
        if not s:
            continue
        matched = False
        for pattern, section_name in _SECTION_PATTERNS:
            if re.match(pattern, s, re.I):
                flush_section()
                current_section = section_name
                current_content = []
                matched = True
                break
        if not matched:
            current_content.append(s)

    flush_section()

    # Build output in priority order: skills first, then summary, experience, projects, education, rest
    # Preamble (contact info, name) is discarded — it has no matching signal
    priority = ["skills", "summary", "experience", "projects", "education", "certifications", "languages"]
    ordered: list[str] = []
    for key in priority:
        if key in sections and sections[key]:
            ordered.append(sections[key])

    return "\n\n".join(ordered) if ordered else text
